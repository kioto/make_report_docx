#!/usr/bin/env python3
# -*- coding:utf-8 -*-

import sys
import pathlib
import json
import docx
from docx.shared import Inches


class ReportDocx(object):
    def __init__(self, jsonfile):
        '''コンストラクタ
        '''
        self.doc = docx.Document()
        if not self.read_json(jsonfile):
            self.doc = None

    def read_json(self, jsonfile):
        '''JSONファイルの読み込み
        '''
        try:
            with open(jsonfile, 'r') as f:
                doc_info = json.load(f)
                self.entry_docx(doc_info)

        except json.JSONDecodeError as e:
            print('JSONDecodeError:', e)
            return False

        return True

    def entry_docx(self, doc_info):
        '''jsonから取り出した情報からdocxに文書を登録
        '''
        try:
            self.entry_title_page(doc_info['title'])
            self.entry_abstract_page(doc_info['abstract'])
            for target in doc_info['reports']:
                self.entry_report_page(target)

        except KeyError as ke:
            print('JSON key error: ', ke)

    def entry_title_page(self, title_info):
        '''表紙をdocxに登録
        '''
        self.doc.add_heading(title_info['title'], level=0)
        self.doc.add_paragraph('\n\n\n')
        self.doc.add_paragraph(u'日付: ' + title_info['date'])
        self.doc.add_paragraph(u'所属: ' + title_info['organization'])
        self.doc.add_paragraph(u'担当: ' + title_info['owner'])
        self.doc.add_page_break()  # 改行

    def entry_abstract_page(self, abst_info):
        '''概要をdocxに登録
        '''
        self.doc.add_heading(abst_info['name'], level=1)
        for line in abst_info['description']:
            self.doc.add_paragraph(line)
        self.doc.add_page_break()  # 改行

    def entry_report_page(self, target_info):
        '''レポートをdocxに登録
        '''
        # ターゲットの説明を登録
        self.doc.add_heading(target_info['heading'], level=1)
        if 'picture' in target_info:
            self.doc.add_picture(target_info['picture'], width=Inches(1.25))
        for line in target_info['description']:
            self.doc.add_paragraph(line)
        self.doc.add_page_break()  # 改行
        # サンプルの登録
        if 'samples' in target_info:
            for sample in target_info['samples']:
                # サンプル名
                self.doc.add_heading(sample['heading'], level=2)
                # 表を作成
                table = self.doc.add_table(rows=1, cols=2)
                # 表(0,0)に画像
                pict_parag = table.rows[0].cells[0].paragraphs[0]
                run = pict_parag.add_run()
                run.add_picture(sample['picture'], height=Inches(1.25))
                # 表(0,1)に説明
                desc = '\n'.join(sample['description'])
                table.rows[0].cells[1].text = desc
                # Memo
                self.doc.add_paragraph('')
                self.doc.add_paragraph('Memo:')
                self.doc.add_paragraph('')
                self.doc.add_paragraph('')

            self.doc.add_page_break()  # 改行

    def save(self, docx_filename):
        '''docxファイルの保存
        '''
        if self.doc:
            self.doc.save(docx_filename)


if __name__ == '__main__':
    if len(sys.argv) != 2:
        print('Usage: python %s <json_file>' % (sys.argv[0]))
        exit()

    jsonfile = pathlib.Path(sys.argv[1])
    if not jsonfile.is_file():
        print('%s: No such file or directory')
        exit()

    repo = ReportDocx(jsonfile)
    repo.save(jsonfile.stem + '.docx')
